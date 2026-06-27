import os
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

doc = Document()

# Define styles
style = doc.styles['Normal']
font = style.font
font.name = 'Arial'
font.size = Pt(11)

# Title
title = doc.add_heading('THIẾT KẾ KHÓA HỌC: ĐÀO TẠO AI ỨNG DỤNG CHO BACK OFFICE', level=0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph('Mục tiêu: Giúp nhân viên thành thạo ứng dụng AI, khắc phục điểm yếu về kỹ năng Prompt, tự động hóa các công việc lặp lại tốn thời gian.\n').alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_heading('I. PHÂN TÍCH NHU CẦU TỪ KHẢO SÁT', level=1)
p = doc.add_paragraph()
p.add_run('Dựa trên kết quả khảo sát từ các bộ phận (HR, Sales, Admin, PQA, Kế toán...), dưới đây là các insight chính:\n').bold = True
doc.add_paragraph('- Các công việc tốn thời gian nhất: Soạn thảo email/tài liệu, tìm kiếm & tổng hợp thông tin, phân tích dữ liệu, review CV ứng viên, lập báo cáo, tổng kết chi phí.', style='List Bullet')
doc.add_paragraph('- Rào cản chính khi dùng AI: Thiếu kỹ năng viết prompt (Prompt Engineering), chưa biết cách áp dụng AI vào công việc cụ thể, lo ngại về tính bảo mật và sự chính xác của AI.', style='List Bullet')
doc.add_paragraph('- Mong muốn sau khóa học: Tiết kiệm thời gian, nâng cao hiệu suất, tự động hóa quy trình (Automation), học kỹ năng Prompt cơ bản và nâng cao.', style='List Bullet')

doc.add_heading('II. LỘ TRÌNH KHÓA HỌC 10 BUỔI', level=1)

sessions = [
    ("Buổi 1: Tổng quan về AI & Tư duy ứng dụng", "Giới thiệu các công cụ AI phổ biến (ChatGPT, Claude, Gemini). Hiểu về hạn chế (Hallucination) và bảo mật dữ liệu công ty."),
    ("Buổi 2: Kỹ thuật Prompt Engineering Cơ bản", "Công thức cấu trúc Prompt chuẩn (Role - Task - Context - Format). Thực hành viết Prompt cơ bản."),
    ("Buổi 3: Prompt Engineering Nâng cao & Chuỗi lệnh (Chain-of-Thought)", "Cách dẫn dắt AI xử lý bài toán phức tạp, cung cấp few-shot (ví dụ mẫu) để kết quả chính xác hơn."),
    ("Buổi 4: Thực hành Demo 1 - Trợ lý giao tiếp & Xử lý văn bản", "Dành cho Sales, Admin, Tư vấn. Hướng dẫn dùng AI để viết/trả lời email, dịch thuật và tinh chỉnh văn phong."),
    ("Buổi 5: Ứng dụng AI trong Đọc hiểu, Tìm kiếm và Tóm tắt tài liệu", "Sử dụng AI (như NotebookLM, Claude) để xử lý lượng lớn tài liệu, quy trình (Quy trình BHXH, Hợp đồng)."),
    ("Buổi 6: Thực hành Demo 2 - Xây dựng Bot tóm tắt và Phân tích", "Dành cho HR, Tuyển dụng. Hướng dẫn dùng AI để review CV, sàng lọc ứng viên và tạo câu hỏi phỏng vấn."),
    ("Buổi 7: Ứng dụng AI trong Phân tích dữ liệu & Báo cáo", "Cách sử dụng AI (Advanced Data Analysis) để xử lý file Excel, vẽ biểu đồ mà không cần dùng hàm phức tạp."),
    ("Buổi 8: Thực hành Demo 3 - Tự động hóa Báo cáo Dữ liệu", "Dành cho PQA, Kế toán, Admin. Tổng hợp lỗi defect, bảng kê hóa đơn và tạo báo cáo trực quan."),
    ("Buổi 9: Sáng tạo Nội dung & Đa phương tiện với AI", "Cách lên Idea, viết content chuẩn SEO, tạo ảnh/video ngắn phục vụ truyền thông nội bộ, bài post Fanpage."),
    ("Buổi 10: Thực hành Demo 4 & Demo 5 - Chuyên đề tự động hóa Workflow & Đánh giá dự án", "Tích hợp AI vào Workflow. Tổng kết khóa học, hướng dẫn các phòng ban xây dựng bộ Prompt chuẩn dùng chung.")
]

table = doc.add_table(rows=1, cols=2)
table.style = 'Table Grid'
hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'Buổi'
hdr_cells[1].text = 'Nội dung chính'
hdr_cells[0].paragraphs[0].runs[0].bold = True
hdr_cells[1].paragraphs[0].runs[0].bold = True

for s in sessions:
    row_cells = table.add_row().cells
    row_cells[0].text = s[0]
    row_cells[1].text = s[1]

doc.add_paragraph('\n')

doc.add_heading('III. YÊU CẦU CHI TIẾT 5 DEMO THỰC HÀNH', level=1)

demos = [
    {
        "name": "Demo 1: Trợ lý Soạn thảo & Trả lời Email Khách hàng (Dành cho Sales, Admin, Tư vấn)",
        "inputs": "Một email khiếu nại hoặc yêu cầu tư vấn của khách hàng. Một file Text/PDF chứa thông tin chính sách của công ty.",
        "process": "Học viên xây dựng một Prompt yêu cầu AI đóng vai chuyên viên CSKH, phân tích cảm xúc của khách, trích xuất chính sách tương ứng và soạn email phản hồi.",
        "outputs": "Email phản hồi lịch sự, đúng văn phong công ty, giải quyết đúng trọng tâm khiếu nại và cung cấp 2-3 tùy chọn/giải pháp cho khách."
    },
    {
        "name": "Demo 2: Hệ thống Lọc và Phân tích CV Ứng viên (Dành cho HR, Tuyển dụng)",
        "inputs": "3 file PDF chứa CV của ứng viên. 1 đoạn mô tả công việc (Job Description - JD).",
        "process": "Học viên dùng tính năng đọc tài liệu của AI, yêu cầu AI đối chiếu kinh nghiệm của từng ứng viên với JD, chấm điểm mức độ phù hợp (thang điểm 10) dựa trên tiêu chí cụ thể.",
        "outputs": "Bảng tóm tắt (Table) so sánh 3 ứng viên gồm các cột: Điểm mạnh, Điểm yếu, Kinh nghiệm tương thích, Điểm đánh giá tổng quan, và 3 Câu hỏi gợi ý cho vòng phỏng vấn với từng ứng viên."
    },
    {
        "name": "Demo 3: Công cụ Phân tích Dữ liệu và Trực quan hóa Báo cáo (Dành cho PQA, Kế toán, Admin)",
        "inputs": "Một file Excel chứa dữ liệu thô (Ví dụ: danh sách defect dự án từ Redmine hoặc Bảng kê chi phí hàng tháng chưa được phân loại).",
        "process": "Học viên upload file dữ liệu lên AI (ChatGPT Data Analyst hoặc Claude), viết prompt yêu cầu dọn dẹp dữ liệu (xử lý giá trị trống), thống kê tần suất và nhóm dữ liệu.",
        "outputs": "Bản báo cáo tóm tắt insight dưới dạng văn bản và 2 biểu đồ trực quan (Biểu đồ tròn tỷ lệ lỗi/chi phí, Biểu đồ cột xu hướng) được AI tự động sinh ra."
    },
    {
        "name": "Demo 4: Trợ lý Sáng tạo Nội dung Đa nền tảng (Dành cho Truyền thông nội bộ, Sales)",
        "inputs": "Một gạch đầu dòng chủ đề cơ bản (Ví dụ: Giới thiệu dịch vụ mới hoặc Sự kiện Teambuilding công ty).",
        "process": "Học viên yêu cầu AI đóng vai Content Creator, triển khai chủ đề gốc thành chuỗi bài viết cho các kênh khác nhau, đảm bảo tính nhất quán nhưng khác biệt về văn phong nền tảng.",
        "outputs": "1 bài đăng Facebook (có emoji, hashtag phù hợp), 1 kịch bản video ngắn TikTok/Shorts (chia rõ kịch bản hình ảnh/âm thanh), và 1 Image Prompt (câu lệnh tiếng Anh) để đưa vào các tool tạo ảnh AI."
    },
    {
        "name": "Demo 5: Trợ lý Phân tích Chân dung Doanh nghiệp & Đề xuất (Dành cho Sales, Tư vấn)",
        "inputs": "Một đường link website hoặc file báo cáo thường niên của một công ty mục tiêu/khách hàng tiềm năng.",
        "process": "Học viên dùng AI có khả năng Browse web hoặc đọc file để trích xuất thông tin lĩnh vực hoạt động, đối thủ cạnh tranh, rào cản của doanh nghiệp đó.",
        "outputs": "Bản phân tích SWOT (Điểm mạnh, Điểm yếu, Cơ hội, Thách thức) của công ty mục tiêu và 1 sườn Đề xuất giải pháp (Proposal) đánh trúng 3 Pain-points để Sales dễ dàng chốt deal."
    }
]

for d in demos:
    doc.add_heading(d["name"], level=2)
    doc.add_paragraph('Đầu vào (Inputs): ', style='List Bullet').runs[0].bold = True
    doc.paragraphs[-1].add_run(d["inputs"])
    
    doc.add_paragraph('Quy trình xử lý (Process): ', style='List Bullet').runs[0].bold = True
    doc.paragraphs[-1].add_run(d["process"])
    
    doc.add_paragraph('Đầu ra (Outputs): ', style='List Bullet').runs[0].bold = True
    doc.paragraphs[-1].add_run(d["outputs"])

doc.save('Khoa_Hoc_AI_VietIS.docx')
